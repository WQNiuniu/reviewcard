from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.opc.constants import RELATIONSHIP_TYPE
import os
import tempfile
import shutil
from lxml import etree

def create_smart_document(input_path, output_path, xml_data_path):
    # 创建临时工作目录
    temp_dir = tempfile.mkdtemp()
    temp_docx = os.path.join(temp_dir, "temp.docx")
    
    # 复制原始文档
    shutil.copy(input_path, temp_docx)
    
    # 打开文档
    doc = Document(temp_docx)
    
    # 添加自定义XML部件
    with open(xml_data_path, 'r', encoding='utf-8') as f:
        xml_content = f.read()
    
    part = doc.part
    custom_xml_part = part._package.create_customxmlpart(
        RELATIONSHIP_TYPE.CUSTOM_XML,
        content_type="application/xml"
    )
    custom_xml_part.blob = xml_content.encode('utf-8')
    
    # 添加XML映射
    namespace = 'http://schemas.openxmlformats.org/officeDocument/2006/customXml'
    schema_ref = '<xs:schema xmlns:xs="http://www.w3.org/2001/XMLSchema">' \
                 '<xs:element name="ProjectData" type="xs:string"/>' \
                 '</xs:schema>'
    
    # 创建内容控件
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                
                # 识别关键字段并替换为内容控件
                if "项目编号" in text:
                    cell.text = ""
                    add_content_control(cell, "项目编号", "ProjectNumber")
                elif "工程名称" in text or "项目名称" in text:
                    cell.text = ""
                    add_content_control(cell, "工程名称", "ProjectName")
                elif "项目经理" in text:
                    cell.text = ""
                    add_content_control(cell, "项目经理", "ProjectManager")
                elif "设计阶段" in text:
                    cell.text = ""
                    add_content_control(cell, "设计阶段", "DesignPhase")
                elif "设计专业" in text:
                    cell.text = ""
                    add_content_control(cell, "设计专业", "Department")
                elif "设计人" in text:
                    cell.text = ""
                    add_content_control(cell, "设计人", "Designer")
    
    # 保存文档
    doc.save(output_path)
    
    # 清理临时目录
    shutil.rmtree(temp_dir)
    print(f"智能文档已创建: {output_path}")

def add_content_control(cell, title, tag):
    """向单元格添加内容控件"""
    # 创建内容控件XML
    sdt_xml = f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{title}"/>
                <w:tag w:val="{tag}"/>
                <w:id w:val="{-abs(hash(tag))}"/>
                <w:placeholder>
                    <w:docPart w:val="DefaultPlaceholder_22675732"/>
                </w:placeholder>
                <w:text/>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:t>{title}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    '''
    
    # 清除单元格现有内容
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    
    # 添加新段落和内容控件
    new_paragraph = cell.add_paragraph()
    sdt_element = parse_xml(sdt_xml)
    new_paragraph._element.append(sdt_element)

def update_project_data(doc_path, field, value):
    """更新项目数据"""
    # 创建临时工作目录
    temp_dir = tempfile.mkdtemp()
    temp_docx = os.path.join(temp_dir, "temp.docx")
    
    # 复制文档
    shutil.copy(doc_path, temp_docx)
    
    # 解压docx文件
    with zipfile.ZipFile(temp_docx, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)
    
    # 查找customXml文件
    custom_xml_files = glob.glob(os.path.join(temp_dir, 'customXml', 'item*.xml'))
    
    if not custom_xml_files:
        print("未找到自定义XML数据")
        shutil.rmtree(temp_dir)
        return
    
    # 更新XML数据
    for xml_file in custom_xml_files:
        tree = etree.parse(xml_file)
        root = tree.getroot()
        
        # 查找并更新字段
        elements = root.xpath(f'//*[local-name()="{field}"]')
        if elements:
            elements[0].text = value
            tree.write(xml_file, encoding="utf-8", xml_declaration=True)
    
    # 重新打包docx文件
    with zipfile.ZipFile(doc_path, 'w') as zipf:
        for root, _, files in os.walk(temp_dir):
            for file in files:
                if file == 'temp.docx':
                    continue
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, temp_dir)
                zipf.write(file_path, arcname)
    
    # 清理临时目录
    shutil.rmtree(temp_dir)
    print(f"项目数据已更新: {field} = {value}")

# 使用示例
if __name__ == "__main__":
    input_file = "校审卡.docx"
    output_file = "智能校审卡.docx"
    xml_data = "project-data.xml"
    
    # 创建智能文档
    create_smart_document(input_file, output_file, xml_data)
    
    # 更新项目数据示例
    update_project_data(output_file, "ProjectManager", "李四")
    update_project_data(output_file, "ProjectNumber", "SCKH-DD20240599")